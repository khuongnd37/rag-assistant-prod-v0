import re
import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path

# Import PyPDF2 với fallback handling
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
    print("✅ PyPDF2 đã được tải thành công")
except ImportError:
    print("⚠️ PyPDF2 chưa được cài đặt. Một số tính năng PDF có thể bị hạn chế.")
    PYPDF2_AVAILABLE = False

# Import python-docx với fallback handling
try:
    import docx
    DOCX_AVAILABLE = True
    print("✅ python-docx đã được tải thành công")
except ImportError:
    print("⚠️ python-docx chưa được cài đặt. Không thể xử lý file DOCX.")
    DOCX_AVAILABLE = False

# Import LangChain với fallback handling
try:
    from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document
    LANGCHAIN_AVAILABLE = True
    print("✅ LangChain đã được tải thành công")
except ImportError:
    print("⚠️ LangChain chưa được cài đặt. Sử dụng chế độ xử lý cơ bản.")
    LANGCHAIN_AVAILABLE = False
    
    # Fallback Document class khi LangChain không khả dụng
    class Document:
        def __init__(self, page_content: str, metadata: Dict[str, Any] = None):
            self.page_content = page_content
            self.metadata = metadata or {}

# Thiết lập logging system
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Xử lý tài liệu toàn diện cho hệ thống RAG
    Hỗ trợ PDF, DOCX, TXT với fallback strategies
    Tối ưu cho file lớn và tích hợp Streamlit
    """
    
    def __init__(self, 
                 chunk_size: int = None, 
                 chunk_overlap: int = None,
                 enable_smart_splitting: bool = True,
                 max_file_size_mb: int = 50):
        """
        Khởi tạo DocumentProcessor với cấu hình tùy chỉnh
        
        Args:
            chunk_size: Kích thước chunk (mặc định 1000 ký tự)
            chunk_overlap: Độ trùng lặp giữa chunks (mặc định 200 ký tự)
            enable_smart_splitting: Bật tách chunk thông minh theo loại tài liệu
            max_file_size_mb: Kích thước file tối đa cho phép (MB)
        """
        self.chunk_size = chunk_size or int(os.getenv("chunk_size", "1000"))
        self.chunk_overlap = chunk_overlap or int(os.getenv("chunk_overlap", "200"))
        self.enable_smart_splitting = enable_smart_splitting
        self.max_file_size_mb = max_file_size_mb
        
        # Patterns nhận diện loại tài liệu để tối ưu xử lý
        self.document_patterns = {
            'legal': [
                r"Chương\s+[IVXLCDM]+", r"Điều\s+\d+", r"Luật\s+\w+",
                r"Nghị định\s+\d+", r"Thông tư\s+\d+", r"Quyết định\s+\d+"
            ],
            'academic': [
                r"Abstract\s*:", r"Tóm tắt\s*:", r"References\s*:",
                r"AI\s+", r"Machine Learning", r"Deep Learning", 
                r"Neural Network", r"Artificial Intelligence",
                r"Methodology\s*:", r"Research\s+", r"Algorithm\s*:",
                r"Data\s+Science", r"Computer\s+Vision", r"NLP"
            ],
            'business': [
                r"Executive Summary", r"Market Analysis", 
                r"Financial Projection", r"Business Plan"
            ]
        }
        
        self._setup_text_splitters()
        logger.info(f"✅ DocumentProcessor khởi tạo thành công: chunk_size={self.chunk_size}, chunk_overlap={self.chunk_overlap}")
    
    def _setup_text_splitters(self):
        """Thiết lập các text splitters chuyên biệt cho từng loại tài liệu"""
        if not LANGCHAIN_AVAILABLE:
            logger.warning("⚠️ LangChain không khả dụng, sử dụng phương pháp xử lý cơ bản")
            return
        
        try:
            # Splitter cho tài liệu học thuật/AI (tối ưu cho file kỹ thuật)
            self.academic_splitter = RecursiveCharacterTextSplitter(
                separators=[
                    r"\n#{1,6}\s",                     # Headers markdown
                    r"\n\d+\.\s",                     # Numbered sections
                    r"\n[A-Z][^\n]*:\s*\n",          # Section headers
                    r"\n\n",                          # Paragraphs
                    r"\n",                            # Lines
                    r"\. ",                           # Sentences
                    r" ",                             # Words
                    r""                               # Characters
                ],
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                length_function=len,
                is_separator_regex=True
            )
            
            # Splitter chuẩn cho tài liệu thông thường
            self.standard_splitter = RecursiveCharacterTextSplitter(
                separators=[r"\n\n", r"\n", r"\. ", r" ", r""],
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                length_function=len
            )
            
            logger.info("✅ Text splitters đã được thiết lập thành công")
            
        except Exception as e:
            logger.error(f"❌ Lỗi thiết lập text splitters: {e}")
            self.academic_splitter = None
            self.standard_splitter = None
    
    def _validate_file(self, file_path: str) -> bool:
        """Kiểm tra tính hợp lệ của file trước khi xử lý"""
        try:
            path = Path(file_path)
            
            if not path.exists():
                logger.error(f"❌ File không tồn tại: {file_path}")
                return False
                
            if not path.is_file():
                logger.error(f"❌ Đường dẫn không phải là file: {file_path}")
                return False
                
            file_size_mb = path.stat().st_size / (1024 * 1024)
            logger.info(f"📄 File: {path.name}")
            logger.info(f"📊 Kích thước: {file_size_mb:.2f} MB")
            logger.info(f"📋 Loại: {path.suffix}")
            
            if file_size_mb > self.max_file_size_mb:
                logger.warning(f"⚠️ File lớn ({file_size_mb:.2f}MB > {self.max_file_size_mb}MB)")
                logger.warning("⚠️ Quá trình xử lý có thể mất nhiều thời gian")
                
            return True
            
        except Exception as e:
            logger.error(f"❌ Lỗi kiểm tra file {file_path}: {e}")
            return False
    
    def detect_document_type(self, content: str) -> str:
        """Phát hiện loại tài liệu dựa trên nội dung để tối ưu xử lý"""
        if not self.enable_smart_splitting:
            return 'standard'
        
        scores = {}
        for doc_type, patterns in self.document_patterns.items():
            score = sum(1 for pattern in patterns 
                       if re.search(pattern, content, re.IGNORECASE | re.MULTILINE))
            scores[doc_type] = score
        
        max_score = max(scores.values()) if scores else 0
        if max_score >= 1:
            detected_type = max(scores, key=scores.get)
            logger.info(f"🔍 Phát hiện loại tài liệu: {detected_type} (điểm: {max_score})")
            return detected_type
        
        logger.info("🔍 Sử dụng loại tài liệu chuẩn")
        return 'standard'
    
    def get_text_splitter(self, doc_type: str):
        """Lấy text splitter phù hợp với loại tài liệu"""
        if not LANGCHAIN_AVAILABLE:
            return None
        
        if doc_type == 'academic' and hasattr(self, 'academic_splitter'):
            return self.academic_splitter
        elif hasattr(self, 'standard_splitter'):
            return self.standard_splitter
        else:
            return None
    
    def _simple_chunk_split(self, text: str) -> List[str]:
        """Tách chunk đơn giản khi LangChain không khả dụng"""
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end < len(text):
                # Tìm điểm cắt tốt nhất để tránh cắt giữa từ/câu
                for separator in ['\n\n', '\n', '. ', ' ']:
                    break_point = text.rfind(separator, start, end)
                    if break_point > start:
                        end = break_point + len(separator)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = max(end - self.chunk_overlap, start + 1)
        
        return chunks
    
    def process_pdf(self, 
                   file_path: str, 
                   use_langchain: bool = True,
                   parent_retriever: bool = True) -> List[Document]:
        """
        Xử lý file PDF với nhiều phương pháp fallback
        
        Args:
            file_path: Đường dẫn file PDF (tham số bắt buộc)
            use_langchain: Ưu tiên sử dụng LangChain (khuyến nghị)
            parent_retriever: Tham số tương thích với code cũ
            
        Returns:
            List[Document]: Danh sách document chunks đã được xử lý
            
        Raises:
            ValueError: Khi file_path trống hoặc file không có nội dung
            FileNotFoundError: Khi file không tồn tại hoặc không hợp lệ
            ImportError: Khi không có thư viện PDF nào được cài đặt
        """
        # Kiểm tra tham số bắt buộc
        if not file_path:
            raise ValueError("file_path là tham số bắt buộc và không được để trống")
        
        if not self._validate_file(file_path):
            raise FileNotFoundError(f"File không hợp lệ hoặc không tồn tại: {file_path}")
        
        try:
            pdf_path = Path(file_path)
            documents = []
            
            logger.info(f"🚀 Bắt đầu xử lý PDF: {pdf_path.name}")
            
            # Phương pháp 1: Sử dụng LangChain (ưu tiên)
            if use_langchain and LANGCHAIN_AVAILABLE:
                logger.info("📚 Sử dụng LangChain PyPDFLoader...")
                loader = PyPDFLoader(str(pdf_path))
                pages = loader.load()
                
                if not pages:
                    raise ValueError(f"Không tìm thấy nội dung trong PDF: {pdf_path}")
                
                logger.info(f"📖 Đã tải {len(pages)} trang từ PDF")
                
                # Phân tích loại tài liệu từ mẫu nội dung
                sample_content = "\n".join([page.page_content[:1000] for page in pages[:5]])
                doc_type = self.detect_document_type(sample_content)
                
                # Tách chunks với splitter phù hợp
                text_splitter = self.get_text_splitter(doc_type)
                if text_splitter:
                    chunks = text_splitter.split_documents(pages)
                    logger.info(f"✂️ Đã tạo {len(chunks)} chunks bằng LangChain splitter")
                else:
                    chunks = pages
                    logger.info(f"✂️ Sử dụng {len(chunks)} chunks từ pages gốc")
                
                # Xử lý và làm sạch chunks
                processed_count = 0
                for idx, chunk in enumerate(chunks):
                    # Bỏ qua chunks quá ngắn (có thể là noise)
                    if len(chunk.page_content.strip()) < 50:
                        continue
                    
                    # Thêm metadata chi tiết cho mỗi chunk
                    source = chunk.metadata.get("source", str(pdf_path))
                    page = chunk.metadata.get("page", 0)
                    
                    chunk.metadata.update({
                        "id": f"{pdf_path.name}:{page}:{idx}",
                        "document_type": doc_type,
                        "chunk_index": idx,
                        "file_name": pdf_path.name,
                        "file_type": "pdf",
                        "processing_method": "langchain",
                        "page_number": page + 1,
                        "chunk_length": len(chunk.page_content)
                    })
                    
                    documents.append(chunk)
                    processed_count += 1
                
                logger.info(f"✅ Đã xử lý {processed_count} chunks hợp lệ")
            
            # Phương pháp 2: Fallback sử dụng PyPDF2
            elif PYPDF2_AVAILABLE:
                logger.info("📚 Sử dụng PyPDF2 fallback...")
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    all_text = ""
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            text = page.extract_text()
                            if text and text.strip():
                                all_text += text.strip() + "\n\n"
                        except Exception as page_error:
                            logger.warning(f"⚠️ Lỗi trích xuất trang {page_num + 1}: {page_error}")
                            continue
                    
                    if not all_text.strip():
                        raise ValueError("Không thể trích xuất text từ PDF")
                    
                    # Phát hiện loại tài liệu
                    doc_type = self.detect_document_type(all_text[:5000])
                    
                    # Tách chunks bằng phương pháp đơn giản
                    chunk_texts = self._simple_chunk_split(all_text)
                    
                    for idx, chunk_text in enumerate(chunk_texts):
                        doc = Document(
                            page_content=chunk_text,
                            metadata={
                                'id': f"{pdf_path.name}:{idx}",
                                'chunk_index': idx,
                                'file_type': 'pdf',
                                'source': str(pdf_path),
                                'file_name': pdf_path.name,
                                'processing_method': 'pypdf2_simple',
                                'document_type': doc_type,
                                'chunk_length': len(chunk_text)
                            }
                        )
                        documents.append(doc)
            
            else:
                # Không có thư viện nào khả dụng
                raise ImportError(
                    "Không có thư viện PDF nào được cài đặt. "
                    "Vui lòng cài đặt một trong các thư viện sau:\n"
                    "- pip install PyPDF2\n"
                    "- pip install langchain-community pypdf"
                )
            
            if not documents:
                raise ValueError(f"Không thể trích xuất nội dung hợp lệ từ PDF: {pdf_path}")
            
            logger.info(f"🎉 Hoàn thành xử lý PDF: {len(documents)} chunks từ {pdf_path.name}")
            return documents
            
        except Exception as e:
            logger.error(f"❌ Lỗi xử lý PDF {file_path}: {str(e)}")
            raise
    
    def process_docx(self, file_path: str, use_langchain: bool = True) -> List[Document]:
        """Xử lý file DOCX với fallback strategies"""
        if not file_path:
            raise ValueError("file_path là tham số bắt buộc và không được để trống")
        
        if not self._validate_file(file_path):
            raise FileNotFoundError(f"File không hợp lệ: {file_path}")
        
        try:
            docx_path = Path(file_path)
            documents = []
            
            logger.info(f"🚀 Bắt đầu xử lý DOCX: {docx_path.name}")
            
            if use_langchain and LANGCHAIN_AVAILABLE:
                # Sử dụng LangChain Docx2txtLoader
                logger.info("📚 Sử dụng LangChain Docx2txtLoader...")
                loader = Docx2txtLoader(str(docx_path))
                pages = loader.load()
                
                if not pages:
                    raise ValueError(f"Không tìm thấy nội dung trong DOCX: {docx_path}")
                
                full_content = pages[0].page_content
                doc_type = self.detect_document_type(full_content[:5000])
                
                text_splitter = self.get_text_splitter(doc_type)
                if text_splitter:
                    chunks = text_splitter.split_documents(pages)
                    logger.info(f"✂️ Đã tạo {len(chunks)} chunks bằng LangChain splitter")
                else:
                    chunks = pages
                    logger.info(f"✂️ Sử dụng {len(chunks)} chunks từ pages gốc")
                
                processed_count = 0
                for idx, chunk in enumerate(chunks):
                    if len(chunk.page_content.strip()) < 50:
                        continue
                    
                    chunk.metadata.update({
                        "id": f"{docx_path.name}:{idx}",
                        "document_type": doc_type,
                        "chunk_index": idx,
                        "file_name": docx_path.name,
                        "file_type": "docx",
                        "processing_method": "langchain",
                        "chunk_length": len(chunk.page_content)
                    })
                    documents.append(chunk)
                    processed_count += 1
                
                logger.info(f"✅ Đã xử lý {processed_count} chunks hợp lệ")
            
            elif DOCX_AVAILABLE:
                # Fallback sử dụng python-docx
                logger.info("📚 Sử dụng python-docx fallback...")
                doc = docx.Document(docx_path)
                
                all_text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        all_text += paragraph.text.strip() + "\n\n"
                
                if not all_text.strip():
                    raise ValueError("Không thể trích xuất text từ DOCX")
                
                doc_type = self.detect_document_type(all_text[:5000])
                chunk_texts = self._simple_chunk_split(all_text)
                
                for idx, chunk_text in enumerate(chunk_texts):
                    if len(chunk_text.strip()) < 50:
                        continue
                    
                    doc_obj = Document(
                        page_content=chunk_text,
                        metadata={
                            'id': f"{docx_path.name}:{idx}",
                            'chunk_index': idx,
                            'file_type': 'docx',
                            'source': str(docx_path),
                            'file_name': docx_path.name,
                            'processing_method': 'python_docx',
                            'document_type': doc_type,
                            'chunk_length': len(chunk_text)
                        }
                    )
                    documents.append(doc_obj)
            
            else:
                raise ImportError(
                    "Không có thư viện DOCX nào được cài đặt. "
                    "Vui lòng cài đặt: pip install python-docx hoặc pip install langchain-community docx2txt"
                )
            
            if not documents:
                raise ValueError(f"Không thể trích xuất nội dung từ DOCX: {docx_path}")
            
            logger.info(f"🎉 Hoàn thành: {len(documents)} chunks từ {docx_path.name}")
            return documents
            
        except Exception as e:
            logger.error(f"❌ Lỗi xử lý DOCX {file_path}: {str(e)}")
            raise
    
    def process_txt(self, file_path: str, encoding: str = 'utf-8', use_langchain: bool = True) -> List[Document]:
        """Xử lý file TXT với multiple encoding support"""
        if not file_path:
            raise ValueError("file_path là tham số bắt buộc và không được để trống")
        
        if not self._validate_file(file_path):
            raise FileNotFoundError(f"File không hợp lệ: {file_path}")
        
        try:
            txt_path = Path(file_path)
            documents = []
            
            logger.info(f"🚀 Bắt đầu xử lý TXT: {txt_path.name}")
            
            if use_langchain and LANGCHAIN_AVAILABLE:
                # Sử dụng LangChain TextLoader với fallback encoding
                logger.info("📚 Sử dụng LangChain TextLoader...")
                
                pages = None
                encodings_to_try = [encoding, 'utf-8', 'cp1252', 'latin-1', 'ascii']
                
                for enc in encodings_to_try:
                    try:
                        loader = TextLoader(str(txt_path), encoding=enc)
                        pages = loader.load()
                        logger.info(f"✅ Đã đọc file với encoding: {enc}")
                        break
                    except UnicodeDecodeError:
                        continue
                    except Exception as e:
                        logger.warning(f"⚠️ Lỗi với encoding {enc}: {e}")
                        continue
                
                if not pages:
                    raise ValueError(f"Không thể đọc file với bất kỳ encoding nào: {txt_path}")
                
                full_content = pages[0].page_content
                doc_type = self.detect_document_type(full_content[:5000])
                
                text_splitter = self.get_text_splitter(doc_type)
                if text_splitter:
                    chunks = text_splitter.split_documents(pages)
                    logger.info(f"✂️ Đã tạo {len(chunks)} chunks bằng LangChain splitter")
                else:
                    chunks = pages
                    logger.info(f"✂️ Sử dụng {len(chunks)} chunks từ pages gốc")
                
                processed_count = 0
                for idx, chunk in enumerate(chunks):
                    if len(chunk.page_content.strip()) < 50:
                        continue
                    
                    chunk.metadata.update({
                        "id": f"{txt_path.name}:{idx}",
                        "document_type": doc_type,
                        "chunk_index": idx,
                        "file_name": txt_path.name,
                        "file_type": "txt",
                        "processing_method": "langchain",
                        "chunk_length": len(chunk.page_content)
                    })
                    documents.append(chunk)
                    processed_count += 1
                
                logger.info(f"✅ Đã xử lý {processed_count} chunks hợp lệ")
            
            else:
                # Fallback đọc file trực tiếp với multiple encoding
                logger.info("📚 Sử dụng đọc file TXT cơ bản...")
                
                all_text = None
                encodings_to_try = [encoding, 'utf-8', 'cp1252', 'latin-1', 'ascii']
                
                for enc in encodings_to_try:
                    try:
                        with open(txt_path, 'r', encoding=enc, errors='ignore') as file:
                            all_text = file.read()
                        logger.info(f"✅ Đã đọc file với encoding: {enc}")
                        break
                    except UnicodeDecodeError:
                        continue
                    except Exception as e:
                        logger.warning(f"⚠️ Lỗi với encoding {enc}: {e}")
                        continue
                
                if not all_text or not all_text.strip():
                    raise ValueError("File TXT trống hoặc không thể đọc")
                
                # Detect document type
                doc_type = self.detect_document_type(all_text[:5000])
                
                # Tách chunks
                chunk_texts = self._simple_chunk_split(all_text)
                
                for idx, chunk_text in enumerate(chunk_texts):
                    if len(chunk_text.strip()) < 50:
                        continue
                    
                    doc = Document(
                        page_content=chunk_text,
                        metadata={
                            'id': f"{txt_path.name}:{idx}",
                            'chunk_index': idx,
                            'file_type': 'txt',
                            'source': str(txt_path),
                            'file_name': txt_path.name,
                            'processing_method': 'basic_txt',
                            'document_type': doc_type,
                            'chunk_length': len(chunk_text)
                        }
                    )
                    documents.append(doc)

            if not documents:
                raise ValueError(f"Không thể trích xuất nội dung từ TXT: {txt_path}")
            
            logger.info(f"🎉 Hoàn thành: {len(documents)} chunks từ {txt_path.name}")
            return documents
            
        except Exception as e:
            logger.error(f"❌ Lỗi xử lý TXT {file_path}: {str(e)}")
            raise
    
    def process_document(self, file_path: str, **kwargs) -> List[Document]:
        """Xử lý tài liệu tự động theo extension của file"""
        if not file_path:
            raise ValueError("file_path là tham số bắt buộc và không được để trống")
        
        file_path_obj = Path(file_path)
        extension = file_path_obj.suffix.lower()
        
        logger.info(f"🔄 Đang xử lý file: {file_path_obj.name} ({extension})")
        
        if extension == '.pdf':
            return self.process_pdf(str(file_path_obj), **kwargs)
        elif extension == '.docx':
            return self.process_docx(str(file_path_obj), **kwargs)
        elif extension == '.txt':
            return self.process_txt(str(file_path_obj), **kwargs)
        else:
            logger.error(f"❌ Không hỗ trợ loại file: {extension}")
            supported_formats = ['.pdf', '.docx', '.txt']
            raise ValueError(f"Loại file '{extension}' không được hỗ trợ. Các định dạng hỗ trợ: {supported_formats}")
    
    def get_statistics(self, documents: List[Document]) -> Dict[str, Any]:
        """Tạo thống kê về các document đã xử lý"""
        if not documents:
            return {
                'total_documents': 0,
                'total_characters': 0,
                'average_chunk_size': 0,
                'file_types': {},
                'document_types': {},
                'min_chunk_size': 0,
                'max_chunk_size': 0,
                'processing_methods': {}
            }
        
        chunk_lengths = [len(doc.page_content) for doc in documents]
        
        stats = {
            'total_documents': len(documents),
            'total_characters': sum(chunk_lengths),
            'average_chunk_size': sum(chunk_lengths) / len(documents),
            'file_types': {},
            'document_types': {},
            'processing_methods': {},
            'min_chunk_size': min(chunk_lengths),
            'max_chunk_size': max(chunk_lengths)
        }
        
        for doc in documents:
            # Đếm file types
            file_type = doc.metadata.get('file_type', 'unknown')
            stats['file_types'][file_type] = stats['file_types'].get(file_type, 0) + 1
            
            # Đếm document types
            doc_type = doc.metadata.get('document_type', 'unknown')
            stats['document_types'][doc_type] = stats['document_types'].get(doc_type, 0) + 1
            
            # Đếm processing methods
            method = doc.metadata.get('processing_method', 'unknown')
            stats['processing_methods'][method] = stats['processing_methods'].get(method, 0) + 1
        
        return stats

# Hàm tiện ích tương thích với code cũ (wrapper function)
def load_and_split_pdf(pdf_path: str, parent_retriver: bool = True) -> List[Document]:
    """
    Hàm wrapper tương thích với code cũ, sử dụng DocumentProcessor để xử lý PDF.
    
    Args:
        pdf_path: Đường dẫn file PDF.
        parent_retriver: Tham số tương thích (không được sử dụng trực tiếp trong logic này).
        
    Returns:
        List[Document]: Danh sách các chunks tài liệu.
        
    Raises:
        ValueError: Khi pdf_path trống
        Exception: Các lỗi khác từ quá trình xử lý
    """
    if not pdf_path:
        raise ValueError("pdf_path không được để trống")
    
    try:
        processor = DocumentProcessor(
            chunk_size=int(os.getenv("chunk_size", "1000")),
            chunk_overlap=int(os.getenv("chunk_overlap", "200")),
            max_file_size_mb=100
        )
        # Truyền rõ ràng tham số file_path
        return processor.process_pdf(
            file_path=pdf_path,
            parent_retriever=parent_retriver
        )
    except Exception as e:
        logger.error(f"❌ Lỗi trong load_and_split_pdf: {e}")
        raise

# Factory function để tạo processor (helper function)
def create_document_processor(chunk_size: int = 1000, 
                            chunk_overlap: int = 200,
                            max_file_size_mb: int = 100) -> DocumentProcessor:
    """
    Factory function để tạo một thể hiện của DocumentProcessor với các tham số mặc định
    hoặc tùy chỉnh.
    
    Args:
        chunk_size: Kích thước chunk mặc định.
        chunk_overlap: Độ trùng lặp chunk mặc định.
        max_file_size_mb: Kích thước file tối đa mặc định.
        
    Returns:
        DocumentProcessor: Một thể hiện mới của DocumentProcessor.
    """
    return DocumentProcessor(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        enable_smart_splitting=True,
        max_file_size_mb=max_file_size_mb
    )

# Hàm tiện ích để xử lý batch files
def process_multiple_documents(file_paths: List[str], 
                             processor: DocumentProcessor = None,
                             **kwargs) -> Dict[str, Any]:
    """
    Xử lý nhiều tài liệu cùng lúc
    
    Args:
        file_paths: Danh sách đường dẫn files
        processor: DocumentProcessor instance (tạo mới nếu None)
        **kwargs: Tham số bổ sung cho processor
        
    Returns:
        Dict chứa kết quả xử lý và thống kê
    """
    if processor is None:
        processor = create_document_processor()
    
    results = {
        'successful': [],
        'failed': [],
        'all_documents': [],
        'statistics': {}
    }
    
    for file_path in file_paths:
        try:
            documents = processor.process_document(file_path, **kwargs)
            results['successful'].append({
                'file_path': file_path,
                'document_count': len(documents),
                'documents': documents
            })
            results['all_documents'].extend(documents)
            logger.info(f"✅ Thành công: {file_path} - {len(documents)} chunks")
            
        except Exception as e:
            results['failed'].append({
                'file_path': file_path,
                'error': str(e)
            })
            logger.error(f"❌ Thất bại: {file_path} - {e}")
    
    # Tạo thống kê tổng hợp
    if results['all_documents']:
        results['statistics'] = processor.get_statistics(results['all_documents'])
    
    return results

# Hàm để kiểm tra khả năng xử lý file
def check_file_support(file_path: str) -> Dict[str, Any]:
    """
    Kiểm tra xem file có thể được xử lý không
    
    Args:
        file_path: Đường dẫn file cần kiểm tra
        
    Returns:
        Dict chứa thông tin về khả năng xử lý file
    """
    path = Path(file_path)
    extension = path.suffix.lower()
    
    support_info = {
        'file_path': file_path,
        'file_name': path.name,
        'extension': extension,
        'exists': path.exists(),
        'is_file': path.is_file() if path.exists() else False,
        'size_mb': 0,
        'supported': False,
        'recommended_method': None,
        'available_methods': []
    }
    
    if support_info['exists'] and support_info['is_file']:
        support_info['size_mb'] = path.stat().st_size / (1024 * 1024)
    
    # Kiểm tra hỗ trợ theo extension
    if extension == '.pdf':
        support_info['supported'] = True
        if LANGCHAIN_AVAILABLE:
            support_info['recommended_method'] = 'langchain'
            support_info['available_methods'].append('langchain')
        if PYPDF2_AVAILABLE:
            support_info['available_methods'].append('pypdf2')
            if not support_info['recommended_method']:
                support_info['recommended_method'] = 'pypdf2'
                
    elif extension == '.docx':
        support_info['supported'] = True
        if LANGCHAIN_AVAILABLE:
            support_info['recommended_method'] = 'langchain'
            support_info['available_methods'].append('langchain')
        if DOCX_AVAILABLE:
            support_info['available_methods'].append('python-docx')
            if not support_info['recommended_method']:
                support_info['recommended_method'] = 'python-docx'
                
    elif extension == '.txt':
        support_info['supported'] = True
        if LANGCHAIN_AVAILABLE:
            support_info['recommended_method'] = 'langchain'
            support_info['available_methods'].append('langchain')
        support_info['available_methods'].append('basic')
        if not support_info['recommended_method']:
            support_info['recommended_method'] = 'basic'
    
    return support_info

# Hàm main để test
def main():
    """Hàm main để test các chức năng của DocumentProcessor"""
    print("🚀 Testing DocumentProcessor...")
    
    # Test khởi tạo
    processor = DocumentProcessor(
        chunk_size=1000,
        chunk_overlap=200,
        max_file_size_mb=50
    )
    print(f"✅ Processor đã khởi tạo: chunk_size={processor.chunk_size}")
    
    # Test file support check
    test_files = ["test.pdf", "test.docx", "test.txt", "test.unknown"]
    print("\n📋 Kiểm tra hỗ trợ file:")
    for file_path in test_files:
        support = check_file_support(file_path)
        status = "✅ Hỗ trợ" if support['supported'] else "❌ Không hỗ trợ"
        print(f"  {file_path}: {status} - {support['available_methods']}")
    
    print("\n🎉 Test hoàn thành!")

if __name__ == "__main__":
    main()
